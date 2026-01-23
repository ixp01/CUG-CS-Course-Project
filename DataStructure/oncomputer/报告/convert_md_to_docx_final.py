#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最终版 Markdown to DOCX Converter
将Markdown文件转换为Word文档，Mermaid代码块转换为格式化文本
"""

import re
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    print("✓ python-docx 已安装")
except ImportError:
    print("❌ 需要安装 python-docx: pip install python-docx")
    sys.exit(1)

class FinalMarkdownToDocxConverter:
    def __init__(self, input_file, output_file):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document()
        self.mermaid_counter = 0
        
        # 设置文档样式
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """设置文档样式"""
        styles = self.doc.styles
        
        # 正文样式 - 宋体五号
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '宋体'
        normal_font.size = Pt(10.5)  # 五号字体
        
        # 创建代码块样式
        try:
            code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(9)
            code_font.color.rgb = RGBColor(64, 64, 64)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.right_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        except:
            pass  # 样式可能已存在
    
    def add_styled_paragraph(self, text, style_name='Normal', bold=False):
        """添加带样式的段落"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        # 根据样式设置字体
        if style_name == 'Heading 1':
            run.font.name = '宋体'
            run.font.size = Pt(16)  # 三号
            run.bold = True
        elif style_name == 'Heading 2':
            run.font.name = '宋体'
            run.font.size = Pt(14)  # 四号
            run.bold = True
        elif style_name == 'Heading 3':
            run.font.name = '宋体'
            run.font.size = Pt(12)  # 小四
            run.bold = True
        elif style_name == 'Code Block':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)  # 五号
            if bold:
                run.bold = True
        
        return para
    
    def add_mermaid_as_text(self, mermaid_code, chart_num):
        """将Mermaid代码添加为格式化文本"""
        # 添加图表标题
        self.add_styled_paragraph(f"图表 {chart_num}：系统架构图", 'Normal', bold=True)
        
        # 添加说明
        self.add_styled_paragraph("（以下为Mermaid图表代码，建议在支持Mermaid的编辑器中查看图形效果）", 'Normal')
        
        # 添加代码块
        lines = mermaid_code.strip().split('\n')
        for line in lines:
            if line.strip():
                self.add_styled_paragraph(f"    {line}", 'Code Block')
            else:
                self.doc.add_paragraph()
        
        # 添加分隔线
        self.add_styled_paragraph("─" * 50, 'Normal')
        self.doc.add_paragraph()
    
    def process_line(self, line):
        """处理单行内容"""
        line = line.rstrip()
        
        if not line.strip():
            # 空行
            self.doc.add_paragraph()
            return
        
        # 处理标题
        if line.startswith('### '):
            # 三级标题
            title_text = line[4:].strip()
            self.add_styled_paragraph(title_text, 'Heading 3')
        elif line.startswith('## '):
            # 二级标题
            title_text = line[3:].strip()
            self.add_styled_paragraph(title_text, 'Heading 2')
        elif line.startswith('# '):
            # 一级标题
            title_text = line[2:].strip()
            self.add_styled_paragraph(title_text, 'Heading 1')
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            # 粗体文本
            text = line[2:-2]
            self.add_styled_paragraph(text, 'Normal', bold=True)
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            list_text = line[2:].strip()
            para = self.doc.add_paragraph(f"• {list_text}")
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
            para.paragraph_format.left_indent = Inches(0.25)
        elif re.match(r'^\d+\.', line):
            # 编号列表
            self.add_styled_paragraph(line, 'Normal')
        elif '---' in line and len(line.strip()) <= 5:
            # 分隔线，添加分页符
            self.doc.add_page_break()
        elif line.startswith('```') and 'cpp' in line:
            # 代码块开始标记
            self.add_styled_paragraph("【C++代码】", 'Normal', bold=True)
        elif line.startswith('```') and line.strip() == '```':
            # 代码块结束标记
            self.doc.add_paragraph()
        elif line.startswith('    ') or line.startswith('\t'):
            # 缩进的代码行
            self.add_styled_paragraph(line, 'Code Block')
        else:
            # 普通段落
            if line.strip():
                self.add_styled_paragraph(line, 'Normal')
    
    def convert(self):
        """执行转换"""
        print(f"📖 开始转换文件: {self.input_file}")
        
        try:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            print(f"❌ 读取文件失败: {str(e)}")
            return False
        
        # 查找所有Mermaid代码块
        mermaid_pattern = r'```mermaid\n(.*?)\n```'
        mermaid_matches = re.findall(mermaid_pattern, content, re.DOTALL)
        
        print(f"🎨 发现 {len(mermaid_matches)} 个Mermaid图表")
        
        # 处理Mermaid代码块
        for i, mermaid_code in enumerate(mermaid_matches):
            placeholder = f"__MERMAID_PLACEHOLDER_{i}__"
            pattern = r'```mermaid\n' + re.escape(mermaid_code) + r'\n```'
            content = re.sub(pattern, placeholder, content, count=1)
        
        # 处理其他代码块
        content = re.sub(r'```cpp\n(.*?)\n```', 
                        lambda m: f"【C++代码】\n{m.group(1)}\n", 
                        content, flags=re.DOTALL)
        
        # 按行处理
        lines = content.split('\n')
        in_code_block = False
        
        for line in lines:
            # 检查Mermaid占位符
            mermaid_match = re.match(r'__MERMAID_PLACEHOLDER_(\d+)__', line.strip())
            if mermaid_match:
                mermaid_index = int(mermaid_match.group(1))
                mermaid_code = mermaid_matches[mermaid_index]
                self.add_mermaid_as_text(mermaid_code, mermaid_index + 1)
                print(f"✅ 已处理Mermaid图表 {mermaid_index + 1}")
                continue
            
            # 检查代码块
            if line.startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block:
                    self.doc.add_paragraph()  # 代码块结束后添加空行
                continue
            
            if in_code_block:
                # 在代码块内，用代码样式
                self.add_styled_paragraph(line, 'Code Block')
            else:
                # 普通内容处理
                self.process_line(line)
        
        # 保存文档
        try:
            self.doc.save(self.output_file)
            print(f"✅ 转换完成！输出文件: {self.output_file}")
            return True
        except Exception as e:
            print(f"❌ 保存文档失败: {str(e)}")
            return False

def main():
    """主函数"""
    print("=" * 60)
    print("📄 最终版 Markdown转Word文档工具")
    print("支持中文字体格式和Mermaid代码展示")
    print("=" * 60)
    
    input_file = "数据结构实习报告_完整版.md"
    output_file = "数据结构实习报告_完整版.docx"
    
    if not os.path.exists(input_file):
        print(f"❌ 输入文件不存在: {input_file}")
        return
    
    print(f"📝 输入文件: {input_file}")
    print(f"📄 输出文件: {output_file}")
    print()
    
    # 执行转换
    converter = FinalMarkdownToDocxConverter(input_file, output_file)
    success = converter.convert()
    
    if success:
        print()
        print("🎉 转换成功完成！")
        print("📋 格式说明:")
        print("  • # 标题 → 宋体三号 (16pt)")
        print("  • ## 标题 → 宋体四号 (14pt)")  
        print("  • ### 标题 → 宋体小四 (12pt)")
        print("  • 正文 → 宋体五号 (10.5pt)")
        print("  • Mermaid图表 → 格式化代码文本")
        print("  • C++代码 → Consolas字体")
        print()
        print("💡 提示：Mermaid图表已转换为代码文本，")
        print("   您可以在支持Mermaid的编辑器中查看图形效果")
    else:
        print()
        print("❌ 转换失败！")

if __name__ == "__main__":
    main() 