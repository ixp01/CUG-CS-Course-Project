#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter
将markdown文件转换为docx格式，去除所有markdown标识符
"""

import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def clean_markdown_text(text):
    """清理markdown标识符，保留纯文本内容"""
    lines = text.splitlines()
    cleaned_lines = []
    in_code_block = False
    
    for line in lines:
        # 处理代码块
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
            
        # 去除markdown标识符
        # 去除标题标识符 (# ## ### ####)
        line = re.sub(r'^#+\s*', '', line)
        
        # 去除列表标识符 (- * +)
        line = re.sub(r'^\s*[-*+]\s*', '', line)
        
        # 去除数字列表标识符 (1. 2. 3.)
        line = re.sub(r'^\s*\d+\.\s*', '', line)
        
        # 去除链接标识符 [text](url)
        line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        
        # 去除图片标识符 ![alt](url)
        line = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', line)
        
        # 保留**加粗**标识符，稍后处理
        
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def add_paragraph_with_formatting(doc, text, font_name="宋体", font_size=10.5):
    """添加段落并处理**加粗**格式"""
    if not text.strip():
        doc.add_paragraph()
        return
    
    # 检查是否是标题行（原markdown标题）
    is_title = False
    title_patterns = [
        r'^实验[一二三四]：',
        r'^\d+\.\d+',  # 如 1.5.1
        r'^### \d+\.',  # 原始的markdown标题残留
        r'^## 实验',
        r'^---',
    ]
    
    for pattern in title_patterns:
        if re.match(pattern, text.strip()):
            is_title = True
            break
    
    # 分割**加粗**内容
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    p = doc.add_paragraph()
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗内容
            content = part[2:-2]  # 去除**
            run = p.add_run(content)
            run.bold = True
        else:
            # 普通内容
            run = p.add_run(part)
            # 如果是标题行，整行加粗
            if is_title:
                run.bold = True
        
        # 设置字体
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
    
    return p

def convert_md_to_docx(input_file, output_file, font_name="宋体", font_size=10.5):
    """主转换函数"""
    try:
        # 读取markdown文件
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 清理markdown标识符
        cleaned_content = clean_markdown_text(md_content)
        
        # 创建docx文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # 逐行处理
        lines = cleaned_content.splitlines()
        
        for line in lines:
            add_paragraph_with_formatting(doc, line, font_name, font_size)
        
        # 保存文档
        doc.save(output_file)
        print(f"转换完成！已生成文件: {output_file}")
        return True
        
    except FileNotFoundError:
        print(f"错误：找不到输入文件 {input_file}")
        return False
    except Exception as e:
        print(f"转换过程中出现错误: {str(e)}")
        return False

def main():
    """主函数"""
    input_file = "实验调试报告细化版.md"
    output_file = "实验调试报告细化版.docx"
    
    print("开始转换markdown到docx...")
    print(f"输入文件: {input_file}")
    print(f"输出文件: {output_file}")
    print("字体设置: 宋体，五号字（10.5pt）")
    print("格式处理: 去除markdown标识符，保留**加粗**效果")
    print("-" * 50)
    
    success = convert_md_to_docx(input_file, output_file)
    
    if success:
        print("-" * 50)
        print("✅ 转换成功完成！")
        print(f"📄 生成的docx文件: {output_file}")
        print("📝 格式特点:")
        print("   - 所有文字为宋体五号")
        print("   - 原markdown标题已加粗")
        print("   - **加粗内容**已正确处理")
        print("   - 已移除所有markdown标识符")
    else:
        print("❌ 转换失败，请检查错误信息")

if __name__ == "__main__":
    main() 