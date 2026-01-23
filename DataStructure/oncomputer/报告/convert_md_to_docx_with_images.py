#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word文档工具 - 支持Mermaid图片转换版
将Mermaid代码转换为真实图片并插入到Word文档中
"""

import os
import re
import base64
import requests
from urllib.parse import quote
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import time

class MermaidImageConverter:
    """Mermaid图片转换器"""
    
    def __init__(self):
        self.services = [
            {
                'name': 'mermaid.ink',
                'url': 'https://mermaid.ink/img/',
                'method': 'base64'
            },
            {
                'name': 'kroki.io',
                'url': 'https://kroki.io/mermaid/svg/',
                'method': 'base64'
            }
        ]
    
    def encode_mermaid(self, mermaid_code):
        """编码Mermaid代码"""
        # 清理代码
        clean_code = mermaid_code.strip()
        # Base64编码
        encoded = base64.b64encode(clean_code.encode('utf-8')).decode('ascii')
        return encoded
    
    def download_image(self, mermaid_code, chart_num):
        """下载Mermaid图片"""
        encoded_code = self.encode_mermaid(mermaid_code)
        
        for service in self.services:
            try:
                print(f"  🌐 尝试使用 {service['name']} 服务...")
                
                if service['name'] == 'mermaid.ink':
                    url = f"{service['url']}{encoded_code}"
                else:
                    url = f"{service['url']}{encoded_code}"
                
                response = requests.get(url, timeout=10)
                
                if response.status_code == 200:
                    print(f"  ✅ 成功获取图片 (大小: {len(response.content)} bytes)")
                    return response.content
                else:
                    print(f"  ❌ 服务返回错误: {response.status_code}")
                    
            except Exception as e:
                print(f"  ❌ 服务 {service['name']} 失败: {str(e)}")
                continue
        
        print(f"  ⚠️  所有服务都失败，将使用文本替代")
        return None

class MarkdownToDocxWithImagesConverter:
    """支持图片的Markdown转Word转换器"""
    
    def __init__(self, input_file, output_file):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document()
        self.image_converter = MermaidImageConverter()
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """设置文档样式"""
        styles = self.doc.styles
        
        # 设置标题样式
        if 'Heading 1' in styles:
            h1_style = styles['Heading 1']
            h1_style.font.name = '宋体'
            h1_style.font.size = Pt(16)  # 三号
            h1_style.font.bold = True
            
        if 'Heading 2' in styles:
            h2_style = styles['Heading 2']
            h2_style.font.name = '宋体'
            h2_style.font.size = Pt(14)  # 四号
            h2_style.font.bold = True
            
        if 'Heading 3' in styles:
            h3_style = styles['Heading 3']
            h3_style.font.name = '宋体'
            h3_style.font.size = Pt(12)  # 小四
            h3_style.font.bold = True
    
    def clean_bold_markers(self, text):
        """清理粗体标记"""
        # 移除 **文本** 格式的粗体标记
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        return text
    
    def add_styled_paragraph(self, text, style_name='Normal', bold=False):
        """添加带样式的段落"""
        # 清理粗体标记
        clean_text = self.clean_bold_markers(text)
        
        para = self.doc.add_paragraph(clean_text)
        
        # 设置段落样式
        if style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
            para.style = style_name
        
        # 设置字体
        for run in para.runs:
            if style_name == 'Heading 1':
                run.font.name = '宋体'
                run.font.size = Pt(16)  # 三号
                run.bold = True
            elif style_name == 'Heading 2':
                run.font.name = '宋体'
                run.font.size = Pt(14)  # 四号
                run.bold = True
            elif style_name == 'Heading 3':
                run.font.name = '宋体'
                run.font.size = Pt(12)  # 小四
                run.bold = True
            elif style_name == 'Code Block':
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            else:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)  # 五号
                if bold:
                    run.bold = True
        
        return para
    
    def add_mermaid_image(self, mermaid_code, chart_num):
        """添加Mermaid图片或文本"""
        print(f"🎨 处理Mermaid图表 {chart_num}...")
        
        # 添加图表标题
        self.add_styled_paragraph(f"图表 {chart_num}：系统架构图", 'Normal', bold=True)
        
        # 尝试下载图片
        image_data = self.image_converter.download_image(mermaid_code, chart_num)
        
        if image_data:
            try:
                # 添加图片到文档
                image_stream = BytesIO(image_data)
                para = self.doc.add_paragraph()
                run = para.runs[0] if para.runs else para.add_run()
                run.add_picture(image_stream, width=Inches(6))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"  ✅ 图片已插入文档")
            except Exception as e:
                print(f"  ❌ 插入图片失败: {str(e)}")
                self.add_mermaid_as_text(mermaid_code, chart_num)
        else:
            # 图片下载失败，使用文本
            self.add_mermaid_as_text(mermaid_code, chart_num)
        
        self.doc.add_paragraph()  # 添加空行
    
    def add_mermaid_as_text(self, mermaid_code, chart_num):
        """将Mermaid代码添加为格式化文本（备用方案）"""
        # 添加说明
        self.add_styled_paragraph("（Mermaid图表代码，建议在支持Mermaid的编辑器中查看图形效果）", 'Normal')
        
        # 添加代码块
        lines = mermaid_code.strip().split('\n')
        for line in lines:
            if line.strip():
                self.add_styled_paragraph(f"    {line}", 'Code Block')
            else:
                self.doc.add_paragraph()
        
        # 添加分隔线
        self.add_styled_paragraph("─" * 50, 'Normal')
    
    def process_line(self, line):
        """处理单行内容"""
        line = line.rstrip()
        
        if not line.strip():
            # 空行
            self.doc.add_paragraph()
            return
        
        # 处理标题
        if line.startswith('### '):
            # 三级标题
            title_text = line[4:].strip()
            self.add_styled_paragraph(title_text, 'Heading 3')
        elif line.startswith('## '):
            # 二级标题
            title_text = line[3:].strip()
            self.add_styled_paragraph(title_text, 'Heading 2')
        elif line.startswith('# '):
            # 一级标题
            title_text = line[2:].strip()
            self.add_styled_paragraph(title_text, 'Heading 1')
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            list_text = line[2:].strip()
            # 清理列表项中的粗体标记
            clean_list_text = self.clean_bold_markers(list_text)
            para = self.doc.add_paragraph(f"• {clean_list_text}")
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
            para.paragraph_format.left_indent = Inches(0.25)
        elif re.match(r'^\d+\.', line):
            # 编号列表
            self.add_styled_paragraph(line, 'Normal')
        elif '---' in line and len(line.strip()) <= 5:
            # 分隔线，添加分页符
            self.doc.add_page_break()
        elif line.startswith('```') and 'cpp' in line:
            # 代码块开始标记
            self.add_styled_paragraph("【C++代码】", 'Normal', bold=True)
        elif line.startswith('```') and line.strip() == '```':
            # 代码块结束标记
            self.doc.add_paragraph()
        elif line.startswith('    ') or line.startswith('\t'):
            # 缩进的代码行
            self.add_styled_paragraph(line, 'Code Block')
        else:
            # 普通段落
            if line.strip():
                self.add_styled_paragraph(line, 'Normal')
    
    def convert(self):
        """执行转换"""
        print(f"📖 开始转换文件: {self.input_file}")
        
        try:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            print(f"❌ 读取文件失败: {str(e)}")
            return False
        
        # 预处理：去除所有粗体标记
        print("🧹 清理粗体标记...")
        content = self.clean_bold_markers(content)
        
        # 查找所有Mermaid代码块
        mermaid_pattern = r'```mermaid\n(.*?)\n```'
        mermaid_matches = re.findall(mermaid_pattern, content, re.DOTALL)
        
        print(f"🎨 发现 {len(mermaid_matches)} 个Mermaid图表")
        
        # 处理Mermaid代码块
        for i, mermaid_code in enumerate(mermaid_matches):
            placeholder = f"__MERMAID_PLACEHOLDER_{i}__"
            pattern = r'```mermaid\n' + re.escape(mermaid_code) + r'\n```'
            content = re.sub(pattern, placeholder, content, count=1)
        
        # 处理其他代码块
        content = re.sub(r'```cpp\n(.*?)\n```', 
                        lambda m: f"【C++代码】\n{m.group(1)}\n", 
                        content, flags=re.DOTALL)
        
        # 按行处理
        lines = content.split('\n')
        in_code_block = False
        
        for line in lines:
            # 检查Mermaid占位符
            mermaid_match = re.match(r'__MERMAID_PLACEHOLDER_(\d+)__', line.strip())
            if mermaid_match:
                mermaid_index = int(mermaid_match.group(1))
                mermaid_code = mermaid_matches[mermaid_index]
                self.add_mermaid_image(mermaid_code, mermaid_index + 1)
                continue
            
            # 检查代码块
            if line.startswith('```'):
                in_code_block = not in_code_block
                if not in_code_block:
                    self.doc.add_paragraph()  # 代码块结束后添加空行
                continue
            
            if in_code_block:
                # 在代码块内，用代码样式
                self.add_styled_paragraph(line, 'Code Block')
            else:
                # 普通内容处理
                self.process_line(line)
        
        # 保存文档
        try:
            self.doc.save(self.output_file)
            print(f"✅ 转换完成！输出文件: {self.output_file}")
            return True
        except Exception as e:
            print(f"❌ 保存文档失败: {str(e)}")
            return False

def main():
    """主函数"""
    print("=" * 60)
    print("🖼️  Markdown转Word文档工具 - 图片增强版")
    print("支持Mermaid图表转换为真实图片")
    print("=" * 60)
    
    input_file = "数据结构实习报告_完整版.md"
    output_file = "数据结构实习报告_图片版.docx"
    
    if not os.path.exists(input_file):
        print(f"❌ 输入文件不存在: {input_file}")
        return
    
    print(f"📝 输入文件: {input_file}")
    print(f"📄 输出文件: {output_file}")
    print()
    
    # 执行转换
    converter = MarkdownToDocxWithImagesConverter(input_file, output_file)
    success = converter.convert()
    
    if success:
        print()
        print("🎉 转换成功完成！")
        print("📋 格式说明:")
        print("  • # 标题 → 宋体三号 (16pt)")
        print("  • ## 标题 → 宋体四号 (14pt)")  
        print("  • ### 标题 → 宋体小四 (12pt)")
        print("  • 正文 → 宋体五号 (10.5pt)")
        print("  • 已清理所有粗体标记 **文本** → 文本")
        print("  • Mermaid图表 → 真实图片 (6英寸宽)")
        print("  • C++代码 → Consolas字体")
        print()
        print("✨ 新功能：")
        print("  • 🖼️  Mermaid代码自动转换为图片")
        print("  • 🌐 支持多个在线服务 (mermaid.ink, kroki.io)")
        print("  • 📐 图片自动居中，6英寸宽度")
        print("  • 🔄 图片失败时自动降级为文本")
    else:
        print()
        print("❌ 转换失败！")

if __name__ == "__main__":
    main() 