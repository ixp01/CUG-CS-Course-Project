import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def clean_markdown(md_text):
    # 1. 去除代码块
    md_text = re.sub(r"```[\\s\\S]*?```", "", md_text)
    # 2. 去除所有#、-等标识符，保留内容
    lines = md_text.splitlines()
    cleaned_lines = []
    for line in lines:
        # 去除#和-开头的行首标识符
        line = re.sub(r"^\\s*#+\\s*", "", line)
        line = re.sub(r"^\\s*-\\s*", "", line)
        # 去除多余空行
        if line.strip() == "":
            cleaned_lines.append("")
        else:
            cleaned_lines.append(line)
    return cleaned_lines

def add_paragraph_with_bold(doc, text, font_name="宋体", font_size=10.5):
    # 处理**加粗**内容
    parts = re.split(r"(\\*\\*.*?\\*\\*)", text)
    p = doc.add_paragraph()
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
    return p

def main():
    input_md = "实验调试报告细化版.md"
    output_docx = "实验调试报告细化版.docx"
    font_name = "宋体"
    font_size = 10.5  # 五号字体

    with open(input_md, "r", encoding="utf-8") as f:
        md_text = f.read()

    lines = clean_markdown(md_text)

    doc = Document()
    for line in lines:
        # 跳过多余空行（可选）
        if line.strip() == "":
            doc.add_paragraph()
            continue
        # 标题行（原markdown标题）加粗
        if re.match(r"^\\s*\\*\\*.*\\*\\*\\s*$", line):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
        else:
            add_paragraph_with_bold(doc, line, font_name, font_size)

    doc.save(output_docx)
    print(f"已生成：{output_docx}")

if __name__ == "__main__":
    main()