#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
改进版 Markdown to DOCX Converter
将Markdown文件转换为Word文档，优化Mermaid图表处理
"""

import re
import os
import sys
import base64
import urllib.parse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    print("✓ python-docx 已安装")
except ImportError:
    print("❌ 需要安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    import requests
    print("✓ requests 已安装")
except ImportError:
    print("❌ 需要安装 requests: pip install requests")
    sys.exit(1)

class ImprovedMarkdownToDocxConverter:
    def __init__(self, input_file, output_file):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document()
        self.mermaid_counter = 0
        
        # 设置文档字体样式
        self.setup_document_styles()
    
    def setup_document_styles(self):
        """设置文档样式"""
        styles = self.doc.styles
        
        # 设置正文样式 - 宋体五号
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '宋体'
        normal_font.size = Pt(10.5)  # 五号 = 10.5pt
        
        # 设置标题1样式 - 宋体三号
        try:
            heading1_style = styles['Heading 1']
            heading1_font = heading1_style.font
            heading1_font.name = '宋体'
            heading1_font.size = Pt(16)  # 三号 = 16pt
        except:
            pass
        
        # 设置标题2样式 - 宋体四号  
        try:
            heading2_style = styles['Heading 2']
            heading2_font = heading2_style.font
            heading2_font.name = '宋体'
            heading2_font.size = Pt(14)  # 四号 = 14pt
        except:
            pass
        
        # 设置标题3样式 - 宋体小四
        try:
            heading3_style = styles['Heading 3']
            heading3_font = heading3_style.font
            heading3_font.name = '宋体'
            heading3_font.size = Pt(12)  # 小四 = 12pt
        except:
            pass
    
    def generate_mermaid_image_url(self, mermaid_code):
        """生成Mermaid图片URL"""
        try:
            # 清理代码
            cleaned_code = mermaid_code.strip()
            
            # 使用kroki.io服务
            encoded_diagram = base64.urlsafe_b64encode(cleaned_code.encode('utf-8')).decode('ascii')
            image_url = f"https://kroki.io/mermaid/png/{encoded_diagram}"
            
            return image_url
        except Exception as e:
            print(f"❌ 生成Mermaid URL失败: {str(e)}")
            return None
    
    def download_mermaid_image(self, mermaid_code):
        """下载Mermaid图表"""
        try:
            self.mermaid_counter += 1
            
            # 先尝试kroki.io
            image_url = self.generate_mermaid_image_url(mermaid_code)
            if image_url:
                print(f"🎨 尝试从kroki.io下载图表 {self.mermaid_counter}")
                response = requests.get(image_url, timeout=15)
                if response.status_code == 200:
                    return response.content
            
            # 备用方案：mermaid.ink
            print(f"🎨 尝试从mermaid.ink下载图表 {self.mermaid_counter}")
            encoded_code = urllib.parse.quote(mermaid_code.strip(), safe='')
            backup_url = f"https://mermaid.ink/img/{encoded_code}"
            response = requests.get(backup_url, timeout=15)
            if response.status_code == 200:
                return response.content
                
            return None
            
        except Exception as e:
            print(f"❌ 下载Mermaid图表失败 {self.mermaid_counter}: {str(e)}")
            return None
    
    def add_styled_paragraph(self, text, style_name='Normal'):
        """添加带样式的段落"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        # 根据样式设置字体
        if style_name == 'Heading 1':
            run.font.name = '宋体'
            run.font.size = Pt(16)  # 三号
            run.bold = True
        elif style_name == 'Heading 2':
            run.font.name = '宋体'
            run.font.size = Pt(14)  # 四号
            run.bold = True
        elif style_name == 'Heading 3':
            run.font.name = '宋体'
            run.font.size = Pt(12)  # 小四
            run.bold = True
        else:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)  # 五号
        
        return para
    
    def process_line(self, line):
        """处理单行内容"""
        line = line.rstrip()
        
        if not line.strip():
            # 空行
            self.doc.add_paragraph()
            return
        
        # 处理标题
        if line.startswith('### '):
            # 三级标题
            title_text = line[4:].strip()
            self.add_styled_paragraph(title_text, 'Heading 3')
        elif line.startswith('## '):
            # 二级标题
            title_text = line[3:].strip()
            self.add_styled_paragraph(title_text, 'Heading 2')
        elif line.startswith('# '):
            # 一级标题
            title_text = line[2:].strip()
            self.add_styled_paragraph(title_text, 'Heading 1')
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            # 粗体文本
            text = line[2:-2]
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            list_text = line[2:].strip()
            para = self.doc.add_paragraph(list_text, style='List Bullet')
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # 编号列表
            list_text = line[3:].strip()
            para = self.doc.add_paragraph(list_text, style='List Number')
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
        elif '---' in line:
            # 分隔线，添加分页符
            self.doc.add_page_break()
        else:
            # 普通段落
            if line.strip():
                self.add_styled_paragraph(line, 'Normal')
    
    def convert(self):
        """执行转换"""
        print(f"📖 开始转换文件: {self.input_file}")
        
        try:
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            print(f"❌ 读取文件失败: {str(e)}")
            return False
        
        # 查找所有Mermaid代码块
        mermaid_pattern = r'```mermaid\n(.*?)\n```'
        mermaid_matches = re.findall(mermaid_pattern, content, re.DOTALL)
        
        print(f"🎨 发现 {len(mermaid_matches)} 个Mermaid图表")
        
        # 预下载所有Mermaid图片
        mermaid_images = {}
        for i, mermaid_code in enumerate(mermaid_matches):
            print(f"⬇️ 预处理Mermaid图表 {i+1}/{len(mermaid_matches)}")
            image_data = self.download_mermaid_image(mermaid_code)
            if image_data:
                mermaid_images[i] = image_data
                print(f"✅ 图表 {i+1} 下载成功")
            else:
                print(f"❌ 图表 {i+1} 下载失败")
        
        # 替换Mermaid代码块为占位符
        for i, mermaid_code in enumerate(mermaid_matches):
            placeholder = f"__MERMAID_PLACEHOLDER_{i}__"
            pattern = r'```mermaid\n' + re.escape(mermaid_code) + r'\n```'
            content = re.sub(pattern, placeholder, content, count=1)
        
        # 处理普通代码块
        content = re.sub(r'```cpp\n(.*?)\n```', r'【C++代码】\n\1', content, flags=re.DOTALL)
        content = re.sub(r'```\n(.*?)\n```', r'【代码块】\n\1', content, flags=re.DOTALL)
        
        # 按行处理
        lines = content.split('\n')
        
        for line in lines:
            # 检查Mermaid占位符
            mermaid_match = re.match(r'__MERMAID_PLACEHOLDER_(\d+)__', line.strip())
            if mermaid_match:
                mermaid_index = int(mermaid_match.group(1))
                
                # 添加图表标题
                self.add_styled_paragraph(f"图表 {mermaid_index + 1}：系统结构图", 'Normal')
                
                if mermaid_index in mermaid_images:
                    # 添加图片
                    para = self.doc.add_paragraph()
                    run = para.add_run()
                    
                    try:
                        # 将图片数据写入临时文件
                        temp_img_path = f"temp_mermaid_{mermaid_index}.png"
                        with open(temp_img_path, 'wb') as img_f:
                            img_f.write(mermaid_images[mermaid_index])
                        
                        # 添加图片到文档
                        run.add_picture(temp_img_path, width=Inches(6))
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # 删除临时文件
                        os.remove(temp_img_path)
                        
                        print(f"✅ 图表 {mermaid_index + 1} 已添加到文档")
                        
                    except Exception as e:
                        print(f"❌ 添加图片失败: {str(e)}")
                        self.add_styled_paragraph(f"【图表 {mermaid_index + 1} 显示失败】", 'Normal')
                else:
                    # 图片下载失败的情况
                    self.add_styled_paragraph(f"【图表 {mermaid_index + 1} - 无法显示】", 'Normal')
                
                # 添加空行
                self.doc.add_paragraph()
                continue
            
            # 处理普通行
            self.process_line(line)
        
        # 保存文档
        try:
            self.doc.save(self.output_file)
            print(f"✅ 转换完成！输出文件: {self.output_file}")
            return True
        except Exception as e:
            print(f"❌ 保存文档失败: {str(e)}")
            return False

def main():
    """主函数"""
    print("=" * 60)
    print("📄 改进版 Markdown转Word文档工具")
    print("支持Mermaid图表转换和中文字体格式")
    print("=" * 60)
    
    input_file = "数据结构实习报告_完整版.md"
    output_file = "数据结构实习报告_完整版.docx"
    
    if not os.path.exists(input_file):
        print(f"❌ 输入文件不存在: {input_file}")
        return
    
    print(f"📝 输入文件: {input_file}")
    print(f"📄 输出文件: {output_file}")
    print()
    
    # 执行转换
    converter = ImprovedMarkdownToDocxConverter(input_file, output_file)
    success = converter.convert()
    
    if success:
        print()
        print("🎉 转换成功完成！")
        print("📋 格式说明:")
        print("  • # 标题 → 宋体三号 (16pt)")
        print("  • ## 标题 → 宋体四号 (14pt)")  
        print("  • ### 标题 → 宋体小四 (12pt)")
        print("  • 正文 → 宋体五号 (10.5pt)")
        print("  • Mermaid图表已转换为图片")
    else:
        print()
        print("❌ 转换失败！")

if __name__ == "__main__":
    main() 