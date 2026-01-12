#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter
将Markdown文件转换为Word文档的Python脚本

支持的功能：
- 标题层级转换
- 表格转换
- 代码块转换
- 列表转换
- 粗体、斜体文本
- Mermaid代码块特殊处理
- 统一设置为宋体5号字体
"""

import re
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

class MarkdownToWordConverter:
    def __init__(self):
        self.document = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """设置文档样式 - 统一使用宋体5号"""
        # 设置默认段落样式为宋体5号
        default_style = self.document.styles['Normal']
        default_style.font.name = '宋体'
        default_style.font.size = Pt(10.5)  # 5号字体约等于10.5pt
        
        # 设置标题样式
        styles = self.document.styles
        
        # 标题1样式 - 宋体5号加粗
        if 'Heading 1' not in styles:
            h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        else:
            h1_style = styles['Heading 1']
        h1_style.font.name = '宋体'
        h1_style.font.size = Pt(10.5)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # 标题2样式 - 宋体5号加粗
        if 'Heading 2' not in styles:
            h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        else:
            h2_style = styles['Heading 2']
        h2_style.font.name = '宋体'
        h2_style.font.size = Pt(10.5)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # 标题3样式 - 宋体5号加粗
        if 'Heading 3' not in styles:
            h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        else:
            h3_style = styles['Heading 3']
        h3_style.font.name = '宋体'
        h3_style.font.size = Pt(10.5)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # 代码样式 - 宋体5号
        if 'Code' not in styles:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        else:
            code_style = styles['Code']
        code_style.font.name = '宋体'
        code_style.font.size = Pt(10.5)
        code_style.font.color.rgb = RGBColor(0, 0, 128)
        
        # 表格样式
        if 'Table' not in styles:
            table_style = styles.add_style('Table', WD_STYLE_TYPE.TABLE)
    
    def convert_markdown_to_word(self, markdown_file, output_file=None):
        """将Markdown文件转换为Word文档"""
        try:
            # 读取Markdown文件
            with open(markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 如果没有指定输出文件，使用默认名称
            if output_file is None:
                output_file = Path(markdown_file).stem + '.docx'
            
            # 处理内容
            self.process_content(content)
            
            # 保存Word文档
            self.document.save(output_file)
            print(f"转换完成！输出文件: {output_file}")
            print("文档已统一设置为宋体5号字体")
            return True
            
        except Exception as e:
            print(f"转换过程中出现错误: {str(e)}")
            return False
    
    def process_content(self, content):
        """处理Markdown内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 处理标题
            if line.startswith('#'):
                self.process_heading(line)
                i += 1
                continue
            
            # 处理表格
            if self.is_table_start(lines, i):
                i = self.process_table(lines, i)
                continue
            
            # 处理代码块
            if line.startswith('```'):
                i = self.process_code_block(lines, i)
                continue
            
            # 处理列表
            if self.is_list_item(line):
                i = self.process_list(lines, i)
                continue
            
            # 处理普通段落
            self.process_paragraph(line)
            i += 1
    
    def process_heading(self, line):
        """处理标题"""
        level = len(line) - len(line.lstrip('#'))
        text = line.lstrip('#').strip()
        
        if level == 1:
            para = self.document.add_heading(text, level=1)
        elif level == 2:
            para = self.document.add_heading(text, level=2)
        elif level == 3:
            para = self.document.add_heading(text, level=3)
        else:
            para = self.document.add_heading(text, level=4)
        
        # 确保标题使用宋体5号
        for run in para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
    
    def is_table_start(self, lines, index):
        """判断是否为表格开始"""
        if index + 1 >= len(lines):
            return False
        
        current_line = lines[index].strip()
        next_line = lines[index + 1].strip()
        
        # 检查是否为表格格式：| 内容 | 内容 |
        if '|' in current_line and '|' in next_line:
            # 检查下一行是否包含分隔符
            if re.match(r'^\|[\s\-:|]+\|$', next_line):
                return True
        return False
    
    def process_table(self, lines, start_index):
        """处理表格"""
        table_lines = []
        i = start_index
        
        # 收集表格行
        while i < len(lines) and '|' in lines[i].strip():
            table_lines.append(lines[i].strip())
            i += 1
        
        if len(table_lines) < 2:
            return start_index + 1
        
        # 解析表格数据
        table_data = []
        for line in table_lines:
            if line.startswith('|') and line.endswith('|'):
                # 移除首尾的 | 并分割
                cells = [cell.strip() for cell in line[1:-1].split('|')]
                table_data.append(cells)
        
        if not table_data:
            return i
        
        # 创建Word表格
        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data)
        
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # 填充表格数据
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_data
                    
                    # 设置所有单元格为宋体5号
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(10.5)
                    
                    # 设置表头样式（第一行）
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        return i
    
    def process_code_block(self, lines, start_index):
        """处理代码块"""
        i = start_index + 1
        code_lines = []
        language = ""
        
        # 检查第一行是否指定了语言
        if i < len(lines) and not lines[i].strip().startswith('```'):
            language = lines[i].strip()
            i += 1
        
        # 收集代码行
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        # 添加代码块标题（如果有语言标识）
        if language:
            title_para = self.document.add_paragraph()
            title_run = title_para.add_run(f"代码块 ({language})")
            title_run.bold = True
            title_run.font.name = '宋体'
            title_run.font.size = Pt(10.5)
        
        # 添加代码内容
        if code_lines:
            # 特殊处理Mermaid代码块
            if language.lower() == 'mermaid':
                self.process_mermaid_block(code_lines)
            else:
                # 普通代码块
                code_para = self.document.add_paragraph()
                code_para.style = 'Code'
                code_text = '\n'.join(code_lines)
                code_run = code_para.add_run(code_text)
                code_run.font.name = '宋体'
                code_run.font.size = Pt(10.5)
        
        return i + 1  # 跳过结束的 ```
    
    def process_mermaid_block(self, code_lines):
        """特殊处理Mermaid代码块"""
        # 添加说明文字
        desc_para = self.document.add_paragraph()
        desc_run = desc_para.add_run("Mermaid流程图代码：")
        desc_run.italic = True
        desc_run.font.name = '宋体'
        desc_run.font.size = Pt(10.5)
        
        # 添加代码
        code_para = self.document.add_paragraph()
        code_para.style = 'Code'
        code_text = '\n'.join(code_lines)
        code_run = code_para.add_run(code_text)
        code_run.font.name = '宋体'
        code_run.font.size = Pt(10.5)
        code_run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色
        
        # 添加说明
        note_para = self.document.add_paragraph()
        note_run = note_para.add_run("注意：此Mermaid代码需要在支持Mermaid的环境中渲染才能显示为流程图。")
        note_run.italic = True
        note_run.font.name = '宋体'
        note_run.font.size = Pt(10.5)
        note_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def is_list_item(self, line):
        """判断是否为列表项"""
        return (line.startswith('- ') or 
                line.startswith('* ') or 
                re.match(r'^\d+\.\s', line))
    
    def process_list(self, lines, start_index):
        """处理列表"""
        i = start_index
        list_items = []
        
        # 收集列表项
        while i < len(lines) and self.is_list_item(lines[i].strip()):
            list_items.append(lines[i].strip())
            i += 1
        
        # 创建列表
        for item in list_items:
            if item.startswith('- ') or item.startswith('* '):
                # 无序列表
                text = item[2:]
                para = self.document.add_paragraph(text, style='List Bullet')
            else:
                # 有序列表
                text = re.sub(r'^\d+\.\s', '', item)
                para = self.document.add_paragraph(text, style='List Number')
            
            # 确保列表项使用宋体5号
            for run in para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
        
        return i
    
    def process_paragraph(self, line):
        """处理普通段落"""
        # 处理粗体和斜体
        line = self.process_text_formatting(line)
        
        para = self.document.add_paragraph(line)
        
        # 确保段落使用宋体5号
        for run in para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
        
        # 如果段落包含特殊格式，需要重新设置
        if '**' in line or '*' in line or '`' in line:
            self.apply_text_formatting(para, line)
    
    def process_text_formatting(self, text):
        """处理文本格式化标记"""
        # 移除Markdown标记，保留纯文本
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 粗体
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # 斜体
        text = re.sub(r'`(.*?)`', r'\1', text)        # 行内代码
        return text
    
    def apply_text_formatting(self, paragraph, original_text):
        """应用文本格式化"""
        # 这里可以实现更复杂的格式化逻辑
        # 目前简化处理，只设置基本样式
        pass

def main():
    parser = argparse.ArgumentParser(description='将Markdown文件转换为Word文档')
    parser.add_argument('input_file', help='输入的Markdown文件路径')
    parser.add_argument('-o', '--output', help='输出的Word文件路径（可选）')
    
    args = parser.parse_args()
    
    # 检查输入文件是否存在
    if not Path(args.input_file).exists():
        print(f"错误：输入文件 '{args.input_file}' 不存在")
        sys.exit(1)
    
    # 创建转换器并执行转换
    converter = MarkdownToWordConverter()
    success = converter.convert_markdown_to_word(args.input_file, args.output)
    
    if success:
        print("转换成功完成！")
    else:
        print("转换失败！")
        sys.exit(1)

if __name__ == "__main__":
    main() 